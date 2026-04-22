"""
Build a single DOCX for a site project: title page, TOC, per-page meta table + content.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models.article import GeneratedArticle
from app.models.blueprint import BlueprintPage
from app.models.project import SiteProject
from app.models.site import Site
from app.models.task import Task
from app.services.html_export import is_html_content, resolve_export_body
from app.services.json_parser import clean_and_parse_json
from app.services.meta_parser import extract_meta_from_parsed, meta_variant_list
from app.services.word_counter import count_content_words

logger = logging.getLogger(__name__)


def _strip_boilerplate_html(soup: BeautifulSoup) -> None:
    for sel in ("script", "style", "nav", "footer", "header"):
        for t in soup.find_all(sel):
            t.decompose()


def _get_all_meta_from_task(task: Task, article: GeneratedArticle | None) -> dict[str, Any]:
    """
    Title/description/H1 via extract_meta_from_parsed; all_variants from results/variants (any case).
    """
    result: dict[str, Any] = {
        "title": "",
        "description": "",
        "h1": "",
        "all_variants": [],
    }

    meta_data: dict[str, Any] | None = None
    if article and article.meta_data and isinstance(article.meta_data, dict):
        meta_data = article.meta_data

    if not meta_data:
        sr = task.step_results or {}
        mg = sr.get("meta_generation", {})
        if isinstance(mg, dict):
            raw = mg.get("result", "")
            if isinstance(raw, str) and raw.strip():
                meta_data = clean_and_parse_json(raw)
            elif isinstance(raw, dict):
                meta_data = raw

    if not meta_data:
        meta_data = {}

    result["all_variants"] = meta_variant_list(meta_data)
    extracted = extract_meta_from_parsed(meta_data)
    result["title"] = extracted["title"]
    result["description"] = extracted["description"]
    result["h1"] = extracted["h1"]

    if article:
        if not result["title"] and article.title:
            result["title"] = article.title
        if not result["description"] and article.description:
            result["description"] = article.description or ""

    return result


def _resolve_single_article_body(article: GeneratedArticle, task: Task | None) -> tuple[str, str]:
    if task:
        body, mode = _get_content_from_task(task, article)
        if body.strip():
            return (body, mode)
    raw = (article.html_content or "").strip() or (article.full_page_html or "").strip()
    if not raw:
        return ("", "plain")
    return (raw, "html" if is_html_content(raw) else "plain")


def _add_simple_article_meta_table(
    doc: Document, keyword: str, word_count: int, title: str, description: str
) -> None:
    rows_data: list[tuple[str, str]] = [
        ("Keyword", keyword or ""),
        ("Word Count", str(int(word_count or 0))),
        ("Title", title or ""),
        ("Description", description or ""),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    try:
        table.style = "Light Shading Accent 1"
    except Exception:
        table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = value
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
        for p in c1.paragraphs:
            p.paragraph_format.space_after = Pt(3)


def build_single_article_docx(article: GeneratedArticle, task: Task | None = None) -> bytes:
    """
    One article → .docx: H1 (or article title) as headline, meta table
    (Keyword, Word Count, Title from meta_generation, Description), then body.
    """
    content, mode = _resolve_single_article_body(article, task)
    if not content.strip():
        raise ValueError("No article content to export")

    display_title = (article.title or "").strip()
    if not display_title and task:
        display_title = (task.main_keyword or "").strip()
    if not display_title:
        display_title = "Article"

    description = (article.description or "").strip()
    headline = display_title
    table_title = display_title
    if task:
        meta = _get_all_meta_from_task(task, article)
        if not description:
            description = str(meta.get("description") or "")
        h1_value = str(meta.get("h1") or "").strip()
        meta_title = str(meta.get("title") or "").strip()
        headline = h1_value or display_title
        table_title = meta_title or display_title

    keyword = (task.main_keyword if task else "") or ""

    wc = article.word_count
    if wc is None or wc == 0:
        wc = count_content_words(content)

    doc = Document()
    t = doc.add_paragraph(headline)
    t.runs[0].font.size = Pt(22)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _add_simple_article_meta_table(doc, keyword, int(wc or 0), table_title, description)
    doc.add_paragraph()

    if mode == "html":
        _html_to_docx_body(doc, content)
    else:
        _add_plain_text_content(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_task_export_docx(db: Session, task: Task) -> bytes:
    """
    Export one task to .docx: prefers saved GeneratedArticle; otherwise
    html_structure > final_editing > primary_generation from step_results.
    """
    article = db.query(GeneratedArticle).filter(GeneratedArticle.task_id == task.id).first()
    if article:
        return build_single_article_docx(article, task)

    content, _src = resolve_export_body(task, None, for_html_export=False)
    if not content.strip():
        raise ValueError("No article or draft content to export")

    meta = _get_all_meta_from_task(task, None)
    synthetic = GeneratedArticle(
        task_id=task.id,
        title=meta.get("title") or task.main_keyword or "Article",
        description=str(meta.get("description") or ""),
        html_content=content,
        word_count=count_content_words(content),
    )
    return build_single_article_docx(synthetic, task)


def _get_content_from_task(task: Task, article: GeneratedArticle | None) -> tuple[str, str]:
    """
    Returns (content, mode) where mode is 'html' or 'plain'.
    Prefer persisted article body and full_page extraction; otherwise same step order
    as pipeline assembly (resolve_export_body), not final_editing-only.
    """
    body, _src = resolve_export_body(task, article, for_html_export=False)
    if not body.strip():
        return ("", "plain")
    return (body, "html" if is_html_content(body) else "plain")


def _merge_runs_into_paragraph(para, element: Tag) -> None:
    """Inline runs inside <p> or similar."""
    for child in element.children:
        if isinstance(child, NavigableString):
            t = str(child)
            if t:
                para.add_run(t)
        elif isinstance(child, Tag):
            name = child.name.lower()
            if name in ("strong", "b"):
                r = para.add_run(child.get_text())
                r.bold = True
            elif name in ("em", "i"):
                r = para.add_run(child.get_text())
                r.italic = True
            elif name == "br":
                para.add_run().add_break()
            elif name == "a":
                href = (child.get("href") or "").strip()
                r = para.add_run(child.get_text())
                r.underline = True
                if href:
                    para.add_run(f" ({href})").italic = True
            elif name == "img":
                alt = (child.get("alt") or "").strip() or "image"
                r = para.add_run(f"[Image: {alt}]")
                r.italic = True
            else:
                _merge_runs_into_paragraph(para, child)


def _add_list_item(doc: Document, text: str, ordered: bool, idx: int) -> None:
    try:
        style = "List Number" if ordered else "List Bullet"
        p = doc.add_paragraph(style=style)
        p.add_run(text)
    except Exception:
        prefix = f"{idx}. " if ordered else "• "
        p = doc.add_paragraph()
        p.add_run(prefix + text)


def _html_table_to_docx(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return
    max_cols = 0
    parsed: list[list[str]] = []
    for tr in rows:
        cells = tr.find_all(["th", "td"])
        parsed.append([c.get_text(" ", strip=True) for c in cells])
        max_cols = max(max_cols, len(cells))
    if max_cols == 0:
        return
    t = doc.add_table(rows=len(parsed), cols=max_cols)
    t.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci in range(max_cols):
            text = row[ci] if ci < len(row) else ""
            t.rows[ri].cells[ci].text = text


def _process_block(doc: Document, node: Tag) -> None:
    name = node.name.lower()

    if name in ("script", "style", "nav", "footer", "header"):
        return

    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(name[1])
        doc.add_heading(node.get_text(" ", strip=True), level=min(level, 9))
        return

    if name == "p":
        p = doc.add_paragraph()
        _merge_runs_into_paragraph(p, node)
        return

    if name == "br":
        doc.add_paragraph().add_run().add_break()
        return

    if name == "ul":
        for li in node.find_all("li", recursive=False):
            _add_list_item(doc, li.get_text(" ", strip=True), False, 0)
        return

    if name == "ol":
        for i, li in enumerate(node.find_all("li", recursive=False), start=1):
            _add_list_item(doc, li.get_text(" ", strip=True), True, i)
        return

    if name == "table":
        _html_table_to_docx(doc, node)
        return

    if name == "img":
        alt = (node.get("alt") or "").strip() or "image"
        p = doc.add_paragraph()
        r = p.add_run(f"[Image: {alt}]")
        r.italic = True
        return

    if name in ("div", "section", "article", "main", "blockquote", "body"):
        for child in node.children:
            if isinstance(child, NavigableString):
                t = str(child).strip()
                if t:
                    doc.add_paragraph(t)
            elif isinstance(child, Tag):
                _process_block(doc, child)
        return

    # Fallback: text only
    txt = node.get_text(" ", strip=True)
    if txt:
        doc.add_paragraph(txt)


def _html_to_docx_body(doc: Document, html: str) -> None:
    soup = BeautifulSoup(html, "html.parser")
    _strip_boilerplate_html(soup)
    root = soup.body if soup.body else soup
    for child in list(root.children):
        if isinstance(child, NavigableString):
            t = str(child).strip()
            if t:
                doc.add_paragraph(t)
        elif isinstance(child, Tag):
            _process_block(doc, child)


def _add_plain_text_content(doc: Document, text: str) -> None:
    parts = text.split("\n\n")
    for para_text in parts:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            p.add_run(line)


def _add_meta_table(
    doc: Document,
    slug: str,
    filename: str,
    meta_info: dict[str, Any],
    main_keyword: str,
    word_count: int,
    additional_keywords: list[str],
) -> None:
    rows_data: list[tuple[str, str]] = [
        ("Slug", slug),
        ("Filename", filename),
        ("Meta Title", str(meta_info.get("title", "") or "")),
        ("Meta Description", str(meta_info.get("description", "") or "")),
        ("H1", str(meta_info.get("h1", "") or "")),
        ("Keyword", main_keyword),
        ("Additional Keywords", ", ".join(additional_keywords)),
        ("Word Count", str(word_count)),
    ]
    all_variants = meta_info.get("all_variants") or []
    if isinstance(all_variants, list) and len(all_variants) > 1:
        for i, variant in enumerate(all_variants):
            if not isinstance(variant, dict):
                continue
            v_title = str(variant.get("Title") or variant.get("title") or "")
            v_desc = str(variant.get("Description") or variant.get("description") or "")
            v_h1 = str(variant.get("H1") or variant.get("h1") or "")
            v_trigger = str(variant.get("Trigger") or variant.get("trigger") or "")
            rows_data.append((f"Variant {i + 1} Title", v_title))
            rows_data.append((f"Variant {i + 1} Description", v_desc))
            rows_data.append((f"Variant {i + 1} H1", v_h1))
            rows_data.append((f"Variant {i + 1} Trigger", v_trigger))
    table = doc.add_table(rows=len(rows_data), cols=2)
    try:
        table.style = "Light Shading Accent 1"
    except Exception:
        table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = value
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
        for p in c1.paragraphs:
            p.paragraph_format.space_after = Pt(3)


def _fmt_slug(slug: str) -> str:
    s = (slug or "").strip()
    if not s:
        return ""
    return s if s.startswith("/") else f"/{s}"


def build_project_docx(db: Session, project_id: str) -> bytes:
    project = db.query(SiteProject).filter(SiteProject.id == project_id).first()
    if not project:
        raise ValueError("Project not found")

    site = db.query(Site).filter(Site.id == project.site_id).first()
    site_label = ""
    if site:
        site_label = f"{site.name} / {site.domain}"

    pages: list[BlueprintPage] = (
        db.query(BlueprintPage)
        .filter(BlueprintPage.blueprint_id == project.blueprint_id)
        .order_by(BlueprintPage.sort_order)
        .all()
    )

    tasks = db.query(Task).filter(Task.project_id == project_id).all()
    by_bp: dict[str, Task] = {}
    for t in tasks:
        if t.blueprint_page_id:
            by_bp[str(t.blueprint_page_id)] = t

    task_ids = [t.id for t in tasks]
    articles_by_task: dict[str, GeneratedArticle] = {}
    if task_ids:
        arts = db.query(GeneratedArticle).filter(GeneratedArticle.task_id.in_(task_ids)).all()
        for a in arts:
            articles_by_task[str(a.task_id)] = a

    pk = project.project_keywords if isinstance(project.project_keywords, dict) else {}
    clustered = pk.get("clustered") if isinstance(pk.get("clustered"), dict) else {}

    doc = Document()
    # Title page
    t = doc.add_paragraph(project.name)
    t.runs[0].font.size = Pt(22)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Seed Keyword: {project.seed_keyword}")
    doc.add_paragraph(f"Site: {site_label}")
    doc.add_paragraph(f"Дата генерации: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Всего страниц (blueprint): {len(pages)}")
    doc.add_paragraph(f"Язык: {project.language} | Страна: {project.country}")

    doc.add_page_break()

    doc.add_heading("Оглавление / Table of Contents", level=1)
    for idx, bp in enumerate(pages, start=1):
        task = by_bp.get(str(bp.id))
        article = articles_by_task.get(str(task.id)) if task else None
        content, _ = _get_content_from_task(task, article) if task else ("", "plain")
        ok = bool(task and task.status == "completed" and content.strip())
        flag = "" if ok else " [не сгенерирована]"
        doc.add_paragraph(f"{idx}. {bp.page_title} (slug: {_fmt_slug(bp.page_slug)}){flag}")

    doc.add_page_break()

    for idx, bp in enumerate(pages, start=1):
        task = by_bp.get(str(bp.id))
        article = articles_by_task.get(str(task.id)) if task else None

        doc.add_heading(f"СТРАНИЦА {idx}: {bp.page_title}", level=1)

        cluster_info = clustered.get(bp.page_slug) if isinstance(clustered, dict) else None
        add_kw: list[str] = []
        if isinstance(cluster_info, dict):
            ak = cluster_info.get("assigned_keywords")
            if isinstance(ak, list):
                add_kw = [str(x).strip() for x in ak if str(x).strip()]

        if not task:
            p = doc.add_paragraph()
            r = p.add_run("[Страница пропущена: нет задачи для этой страницы blueprint.]")
            r.italic = True
            doc.add_page_break()
            continue

        meta_info = _get_all_meta_from_task(task, article)
        content, mode = _get_content_from_task(task, article)

        if task.status != "completed" or not content.strip():
            logger.warning(
                "docx export: skip empty content for page %s task=%s status=%s",
                bp.page_slug,
                task.id,
                task.status,
            )
            p = doc.add_paragraph()
            r = p.add_run("[Нет завершённого контента для этой страницы — пропуск основного текста.]")
            r.italic = True
            wc = count_content_words(content) if content else 0
            _add_meta_table(
                doc,
                _fmt_slug(bp.page_slug),
                bp.filename or "",
                meta_info,
                task.main_keyword or "",
                wc,
                add_kw,
            )
            doc.add_page_break()
            continue

        wc = article.word_count if article and article.word_count else count_content_words(content)

        _add_meta_table(
            doc,
            _fmt_slug(bp.page_slug),
            bp.filename or "",
            meta_info,
            task.main_keyword or "",
            int(wc or 0),
            add_kw,
        )
        doc.add_paragraph()

        if mode == "html":
            _html_to_docx_body(doc, content)
        else:
            _add_plain_text_content(doc, content)

        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
